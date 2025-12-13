"""
Script to convert API_ENDPOINTS_DOCUMENTATION.md to DOCX format
Requires: pip install python-docx markdown
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def read_markdown(file_path):
    """Read markdown file content"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def parse_markdown_to_docx(md_content, output_path):
    """Convert markdown to DOCX"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Title (single #)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            i += 1
            continue
        
        # Heading level 1 (##)
        if line.startswith('## '):
            title = line[3:].strip()
            p = doc.add_heading(title, level=1)
            i += 1
            continue
        
        # Heading level 2 (###)
        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_heading(title, level=2)
            i += 1
            continue
        
        # Heading level 3 (####)
        if line.startswith('#### '):
            title = line[5:].strip()
            p = doc.add_heading(title, level=3)
            i += 1
            continue
        
        # Code block
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            i += 1
            continue
        
        # Bullet point
        if line.startswith('- '):
            text = line[2:].strip()
            # Check for bold
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Regular paragraph
        if line:
            # Handle inline code
            text = line
            p = doc.add_paragraph()
            
            # Split by code blocks
            parts = re.split(r'(`[^`]+`)', text)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    # Code
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    # Regular text
                    # Handle bold
                    bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
                    for bold_part in bold_parts:
                        if bold_part.startswith('**') and bold_part.endswith('**'):
                            run = p.add_run(bold_part[2:-2])
                            run.bold = True
                        else:
                            run = p.add_run(bold_part)
        
        i += 1
    
    # Save document
    doc.save(output_path)
    print(f"✅ Đã tạo file DOCX: {output_path}")

if __name__ == '__main__':
    input_file = 'API_ENDPOINTS_DOCUMENTATION.md'
    output_file = 'API_ENDPOINTS_DOCUMENTATION.docx'
    
    try:
        md_content = read_markdown(input_file)
        parse_markdown_to_docx(md_content, output_file)
        print(f"\n📄 File đã được tạo thành công!")
        print(f"📁 Vị trí: {output_file}")
    except FileNotFoundError:
        print(f"❌ Không tìm thấy file: {input_file}")
    except Exception as e:
        print(f"❌ Lỗi: {e}")
        print("\n💡 Hãy cài đặt thư viện cần thiết:")
        print("   pip install python-docx")

